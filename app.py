import streamlit as st
import pandas as pd
from datetime import datetime, timedelta, date
import json
import sys
import os
import io
import re
from docx import Document
from io import BytesIO 

# Добавляем папки в путь
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

from database.models import Database
from utils.calendar_helper import CalendarHelper
from utils.ktp_helper import KTPHelper, validate_ktp_structure
from utils.excel_helper import ExcelHelper
from utils.notes_calendar import NotesCalendar
from utils.ktp_sor_soch import read_ktp_and_extract_sor_soch, create_excel_schedule


# Попытка импорта голосовых библиотек
try:
    import speech_recognition as sr
    VOICE_AVAILABLE = True
except:
    VOICE_AVAILABLE = False

try:
    import pyttsx3
    TTS_AVAILABLE = True
except:
    TTS_AVAILABLE = False

# Инициализация
db = Database()
calendar_helper = CalendarHelper()
ktp_helper = KTPHelper()
excel_helper = ExcelHelper()
notes_calendar = NotesCalendar()

# ===== ИНИЦИАЛИЗАЦИЯ SESSION_STATE =====
if 'user' not in st.session_state:
    st.session_state.user = None
if 'current_page' not in st.session_state:
    st.session_state.current_page = 'login'
if 'chat_mode' not in st.session_state:
    st.session_state.chat_mode = 'text'
if 'chat_history' not in st.session_state:
    st.session_state.chat_history = []
if 'selected_microphone' not in st.session_state:
    st.session_state.selected_microphone = None
if 'voice_note_text' not in st.session_state:
    st.session_state.voice_note_text = ''

# ===== ФУНКЦИИ ОБРАБОТКИ ЗАПРОСОВ =====

def process_calendar_query(query):
    """Обработка запросов про календарь"""
    today = date.today()
    query_lower = query.lower()
    
    current_quarter = calendar_helper.get_current_quarter()
    
    if current_quarter and isinstance(current_quarter, int) and current_quarter > 0:
        info = calendar_helper.get_quarter_info(current_quarter)
        if info:
            if info['days_until_start'] > 0:
                return f"""⏳ {info['quarter']}-я четверть еще не началась:
• Начало: {info['start_date'].strftime('%d.%m.%Y')}
• До начала: {info['days_until_start']} дней
• Рабочих дней: {info['working_days_count']}"""
            elif info['days_until_end'] >= 0:
                return f"""📖 {info['quarter']}-я четверть идет:
• Начало: {info['start_date'].strftime('%d.%m.%Y')}
• Конец: {info['end_date'].strftime('%d.%m.%Y')}
• До конца: {info['days_until_end']} дней
• Осталось рабочих дней: {info['working_days_left']}"""
    
    return "Информация о четверти недоступна"


def process_teacher_query(query, user):
    """
    Обработка запросов учителя.
    Отвечает на информационные вопросы и создает кнопки для навигации.
    """
    assistant_name = user.get('assistant_name', 'Помощник')
    user_name = user.get('user_nickname', '')
    query_lower = query.lower().strip()
    
    # ===== ОПРЕДЕЛЕНИЕ ТЕКУЩЕЙ ДАТЫ И ПЕРИОДА =====
    
    today = date.today()
    today_str = today.strftime('%d.%m.%Y')
    
    # Определяем период
    current_quarter = calendar_helper.get_current_quarter()
    period_info = ""
    
    if current_quarter == 0:
        period_info = "☀️ Летние каникулы"
    elif current_quarter and isinstance(current_quarter, str) and 'каникулы' in current_quarter:
        q_num = current_quarter.split('_')[1]
        period_info = f"🏖️ Каникулы (после {q_num} четверти)"
    elif current_quarter and isinstance(current_quarter, int):
        period_info = f"📖 {current_quarter} четверть"
    else:
        period_info = "🏖️ Каникулы"
    
    # ===== ВОПРОСЫ ПРО ДАТУ И ПЕРИОД =====
    
    if any(word in query_lower for word in ['какая дата', 'какое число', 'сегодня', 'дата сегодня']):
        response = f"""
📅 **Сегодня: {today_str}**

Сейчас: **{period_info}**

"""
        if current_quarter and isinstance(current_quarter, int):
            info = calendar_helper.get_quarter_info(current_quarter)
            if info:
                response += f"""
• Начало четверти: {info['start_date'].strftime('%d.%m.%Y')}
• Конец четверти: {info['end_date'].strftime('%d.%m.%Y')}
• До конца: {info['days_until_end']} дней ({info['weeks_until_end']} недель)
• Рабочих дней осталось: {info['working_days_left']}
"""
        return response, None
    
    # ===== НАВИГАЦИЯ ЧЕРЕЗ КНОПКИ =====
    
    # Создание КТП
    if any(word in query_lower for word in ['создать ктп', 'создай ктп', 'шаблон ктп', 'новый ктп']):
        st.session_state.navigation_request = 'ktp_create'
        return f"""
📚 **Создание шаблона КТП**

Нажмите кнопку ниже, чтобы перейти:
""", 'button_ktp_create'
    
    # Заполнение дат в КТП
    if any(word in query_lower for word in ['заполнить дат', 'даты в ктп', 'заполни дат', 'проставить дат']):
        st.session_state.navigation_request = 'ktp_fill_dates'
        return f"""
📅 **Заполнение дат в КТП**

Нажмите кнопку ниже, чтобы перейти:
""", 'button_ktp_fill'
    
    # График СОР/СОЧ
    if any(word in query_lower for word in ['сор', 'соч', 'суммативн', 'график сор']):
        st.session_state.navigation_request = 'ktp_sor_soch'
        return f"""
📊 **График СОР/СОЧ**

Нажмите кнопку ниже, чтобы перейти:
""", 'button_ktp_sor'
    
    # Список класса
    if any(word in query_lower for word in ['список класса', 'найди ученика', 'ученик', 'адрес ученика']):
        st.session_state.navigation_request = 'class_list'
        return f"""
👥 **Список класса**

Нажмите кнопку ниже, чтобы перейти:
""", 'button_class_list'
    
    # Заметки
    if any(word in query_lower for word in ['создать заметк', 'заметк', 'напомин', 'запиши']):
        st.session_state.navigation_request = 'notes'
        return f"""
📝 **Заметки**

Нажмите кнопку ниже, чтобы перейти:
""", 'button_notes'
    
    # Календарь
    if any(word in query_lower for word in ['календар', 'четверт', 'каникул', 'праздник']):
        st.session_state.navigation_request = 'calendar'
        return f"""
📅 **Календарь**

Нажмите кнопку ниже, чтобы перейти:
""", 'button_calendar'
    
    # ===== ИНФОРМАЦИОННЫЕ ОТВЕТЫ =====
    
    # Сколько дней до конца четверти
    if 'сколько' in query_lower and ('дней' in query_lower or 'недель' in query_lower) and 'четверт' in query_lower:
        if current_quarter and isinstance(current_quarter, int):
            info = calendar_helper.get_quarter_info(current_quarter)
            return f"""
⏰ До конца **{current_quarter} четверти**:

• Дней: **{info['days_until_end']}**
• Недель: **{info['weeks_until_end']}**
• Рабочих дней: **{info['working_days_left']}**

Конец: {info['end_date'].strftime('%d.%m.%Y')}
""", None
    
    # Когда каникулы
    if 'когда' in query_lower and 'каникул' in query_lower:
        if current_quarter and isinstance(current_quarter, int):
            hol_start, hol_end = calendar_helper.get_quarter_holidays(current_quarter)
            if hol_start:
                days_left = (hol_start - today).days
                return f"""
🏖️ **Каникулы после {current_quarter} четверти:**

• Начало: {hol_start.strftime('%d.%m.%Y')}
• Конец: {hol_end.strftime('%d.%m.%Y')}
• Через: {days_left} дней
""", None
    
    # Заметки на сегодня
    if 'заметк' in query_lower and 'сегодня' in query_lower:
        notes = db.get_notes_by_date(user['id'], today.strftime('%Y-%m-%d'))
        if notes:
            response = f"📝 **Заметки на сегодня ({today_str}):**\n\n"
            for note in notes:
                emoji = {'high': '🔴', 'normal': '🟠', 'low': '🟢'}.get(note.get('priority', 'normal'), '⚪')
                response += f"{emoji} {note['title']}\n"
            return response, None
        else:
            return f"📝 На сегодня ({today_str}) заметок нет.", None
    
    # ===== ПРИВЕТСТВИЯ =====
    
    if any(word in query_lower for word in ['привет', 'здравств', 'hi', 'hello']):
        return f"""
👋 Здравствуйте, **{user_name}**!

Сегодня: **{today_str}**
Сейчас: **{period_info}**

**Спросите меня:**
• "Сколько дней до конца четверти?"
• "Какие заметки на сегодня?"
• "Когда каникулы?"

**Или для работы с файлами:**
• "Создать КТП"
• "Заполнить даты в КТП"
• "График СОР/СОЧ"
""", None
    
    # ===== ЕСЛИ НЕ РАСПОЗНАНО =====
    
    return f"""
😊 Извините, я еще не умею этого делать, но обязательно научусь!

**Попробуйте:**
• "Какая сегодня дата?"
• "Сколько дней до конца четверти?"
• "Создать КТП"
• "Заполнить даты в КТП"
• "График СОР/СОЧ"
• "Список класса"
""", None


def process_calendar_info(query):
    """
    Отвечает на вопросы про календарь прямо в чате
    """
    query_lower = query.lower()
    today = date.today()
    
    # Текущая четверть
    current_quarter = calendar_helper.get_current_quarter()
    
    # "Какая сейчас четверть?"
    if 'какая' in query_lower and 'четверть' in query_lower:
        if current_quarter == 0:
            return """
☀️ Сейчас **летние каникулы**!

Новый учебный год начнется **1 сентября 2026 года**.
"""
        elif current_quarter and isinstance(current_quarter, int):
            info = calendar_helper.get_quarter_info(current_quarter)
            if info:
                return f"""
📅 Сейчас идет **{current_quarter} четверть**.

**Даты:**
• Начало: {info['start_date'].strftime('%d.%m.%Y')}
• Конец: {info['end_date'].strftime('%d.%m.%Y')}

**Осталось:**
• Дней: {info['days_until_end']}
• Недель: {info['weeks_until_end']}
• Рабочих дней: {info['working_days_left']}
"""
        elif current_quarter and isinstance(current_quarter, str) and 'каникулы' in current_quarter:
            q_num = current_quarter.split('_')[1]
            return f"""
🏖️ Сейчас **каникулы** (после {q_num} четверти).

Следующая четверть скоро начнется!
"""
    
    # "Сколько дней до конца четверти?"
    if 'сколько' in query_lower and ('дней' in query_lower or 'недель' in query_lower) and 'четверть' in query_lower:
        if current_quarter and isinstance(current_quarter, int):
            info = calendar_helper.get_quarter_info(current_quarter)
            if info:
                return f"""
⏰ До конца **{current_quarter} четверти**:

• 📅 Дней: **{info['days_until_end']}**
• 📆 Недель: **{info['weeks_until_end']}**
• 💼 Рабочих дней: **{info['working_days_left']}**

Конец четверти: {info['end_date'].strftime('%d.%m.%Y')}
"""
    
    # "Когда каникулы?"
    if 'когда' in query_lower and 'каникул' in query_lower:
        if current_quarter and isinstance(current_quarter, int):
            hol_start, hol_end = calendar_helper.get_quarter_holidays(current_quarter)
            if hol_start and hol_end:
                days_left = (hol_start - today).days
                return f"""
🏖️ **Ближайшие каникулы:**

• После: {current_quarter} четверти
• Начало: **{hol_start.strftime('%d.%m.%Y')}**
• Конец: **{hol_end.strftime('%d.%m.%Y')}**
• Через: **{days_left} дней**

Длительность: {(hol_end - hol_start).days + 1} дней
"""
    
    # "Какие праздники?"
    if 'праздник' in query_lower:
        upcoming = calendar_helper.get_next_holidays(5)
        if upcoming:
            response = "🎉 **Ближайшие праздники и каникулы:**\n\n"
            for event in upcoming:
                days_left = event['days_left']
                if days_left == 0:
                    days_text = "🎉 Сегодня!"
                elif days_left == 1:
                    days_text = "Завтра!"
                else:
                    days_text = f"через {days_left} дн."
                
                response += f"• {event['date'].strftime('%d.%m.%Y')} — {event['name']} ({days_text})\n"
            return response
    
    # Общая информация
    if current_quarter and isinstance(current_quarter, int):
        info = calendar_helper.get_quarter_info(current_quarter)
        return f"""
📅 **{current_quarter} четверть**

• Начало: {info['start_date'].strftime('%d.%m.%Y')}
• Конец: {info['end_date'].strftime('%d.%m.%Y')}
• До конца: {info['days_until_end']} дней

Спросите: "Когда каникулы?" или "Какие праздники?"
"""
    
    return "Информация о четверти недоступна."


def process_notes_info(query, user):
    """
    Отвечает на вопросы про заметки прямо в чате
    """
    query_lower = query.lower()
    today = date.today()
    
    # "Какие заметки на сегодня?"
    if 'сегодня' in query_lower:
        notes = db.get_notes_by_date(user['id'], today.strftime('%Y-%m-%d'))
        
        if notes:
            response = f"📝 **Заметки на сегодня ({today.strftime('%d.%m.%Y')}):**\n\n"
            for note in notes:
                priority_emoji = {'high': '🔴', 'normal': '🟠', 'low': '🟢'}.get(note.get('priority', 'normal'), '⚪')
                title = note.get('title', 'Без названия')
                content = note.get('content', '')
                response += f"{priority_emoji} **{title}**\n"
                if content:
                    response += f"   {content[:100]}\n"
                response += "\n"
            return response
        else:
            return f"""
📝 На сегодня ({today.strftime('%d.%m.%Y')}) заметок нет.

Хотите создать? Перейдите в раздел **"Заметки"**.
"""
    
    # "Есть ли заметки на неделю?"
    if 'недел' in query_lower or 'ближайш' in query_lower:
        upcoming = db.get_upcoming_reminders(user['id'], 7)
        
        if upcoming:
            response = "📝 **Ближайшие заметки (на неделю):**\n\n"
            for note in upcoming[:10]:
                priority_emoji = {'high': '🔴', 'normal': '🟠', 'low': '🟢'}.get(note.get('priority', 'normal'), '⚪')
                title = note.get('title', '')
                reminder = note.get('reminder_date', '')
                if reminder:
                    # Форматируем дату
                    try:
                        dt = datetime.strptime(reminder[:10], '%Y-%m-%d')
                        date_str = dt.strftime('%d.%m.%Y')
                    except:
                        date_str = reminder[:10]
                    response += f"{priority_emoji} **{title}** — {date_str}\n"
            return response
        else:
            return "📝 На ближайшую неделю заметок нет."
    
    # "Сколько заметок?"
    if 'сколько' in query_lower:
        notes = db.get_user_notes(user['id'])
        return f"📝 У вас всего **{len(notes)}** заметок."
    
    # "Есть ли напоминания?"
    if 'напомин' in query_lower:
        upcoming = db.get_upcoming_reminders(user['id'], 3)
        if upcoming:
            return f"🔔 У вас **{len(upcoming)}** напоминаний в ближайшие 3 дня."
        else:
            return "🔔 Напоминаний в ближайшие 3 дня нет."
    
    # Если просят создать заметку — направляем
    return f"""
📝 Для создания заметки перейдите в раздел **"Заметки"**.

Там вы можете:
• Создать текстовую заметку
• Создать голосовую заметку
• Установить напоминание

*Переключаю вас...*
""", None

# Функции для работы с голосом
def list_microphones():
    """Получить список доступных микрофонов"""
    if not VOICE_AVAILABLE:
        return []
    
    try:
        import speech_recognition as sr
        return sr.Microphone.list_microphone_names()
    except:
        return []

def voice_to_text(device_index=None):
    """Распознавание речи с выбором микрофона"""
    if not VOICE_AVAILABLE:
        st.error("Модуль распознавания речи не установлен")
        return None
    
    import speech_recognition as sr
    
    recognizer = sr.Recognizer()
    
    # Настройка для лучшего распознавания
    recognizer.energy_threshold = 300  # Чувствительность микрофона
    recognizer.dynamic_energy_threshold = True
    recognizer.pause_threshold = 0.8  # Пауза между словами
    
    try:
        # Если указан конкретный микрофон
        if device_index is not None:
            microphone = sr.Microphone(device_index=device_index)
        else:
            # Используем микрофон по умолчанию
            microphone = sr.Microphone()
        
        st.info("🎤 Говорите... (подождите 1-2 секунды)")
        
        with microphone as source:
            # Калибровка микрофона
            recognizer.adjust_for_ambient_noise(source, duration=1)
            st.info("🎤 Слушаю... Говорите!")
            
            try:
                audio = recognizer.listen(source, timeout=10, phrase_time_limit=15)
                st.info("⏳ Обрабатываю речь...")
            except sr.WaitTimeoutError:
                st.warning("⏰ Время ожидания истекло. Попробуйте еще раз.")
                return None
        
        # Распознавание речи
        try:
            # Пробуем Google (нужен интернет)
            text = recognizer.recognize_google(audio, language='ru-RU')
            return text
        except sr.UnknownValueError:
            st.error("❌ Не удалось распознать речь. Попробуйте говорить четче.")
            return None
        except sr.RequestError:
            # Если нет интернета, пробуем офлайн распознавание
            try:
                # Попытка использовать Vosk (офлайн)
                text = recognizer.recognize_vosk(audio, language='ru')
                return text
            except:
                st.error("❌ Ошибка распознавания. Проверьте интернет-соединение.")
                return None
    
    except Exception as e:
        st.error(f"❌ Ошибка микрофона: {str(e)}")
        st.info("💡 Попробуйте выбрать другой микрофон в настройках")
        return None

def text_to_speech(text):
    """Озвучивание текста"""
    if not TTS_AVAILABLE:
        return
    
    try:
        engine = pyttsx3.init()
        # Настройка русского голоса
        voices = engine.getProperty('voices')
        for voice in voices:
            if 'russian' in voice.name.lower():
                engine.setProperty('voice', voice.id)
                break
        
        engine.say(text)
        engine.runAndWait()
    except:
        pass

# СТРАНИЦЫ ПРИЛОЖЕНИЯ

def login_page():
    """Страница входа"""
    st.markdown('<div class="main-header"><h1>📚 Умный помощник учителя</h1></div>', 
                unsafe_allow_html=True)
    
    col1, col2, col3 = st.columns([1, 2, 1])
    
    with col2:
        tab1, tab2 = st.tabs(["🔑 Вход", "📝 Регистрация"])
        
        with tab1:
            with st.form("login_form"):
                username = st.text_input("👤 Логин")
                password = st.text_input("🔒 Пароль", type="password")
                
                if st.form_submit_button("🚀 Войти", use_container_width=True):
                    user = db.authenticate_user(username, password)
                    if user:
                        st.session_state.user = user
                        st.session_state.current_page = 'main'
                        st.session_state.chat_history = db.get_chat_history(user['id'])
                        st.success(f"✅ Добро пожаловать, {user.get('user_nickname', username)}!")
                        st.rerun()
                    else:
                        st.error("❌ Неверный логин или пароль")
        
        with tab2:
            with st.form("register_form"):
                username = st.text_input("👤 Придумайте логин *")
                password = st.text_input("🔒 Придумайте пароль *", type="password")
                password2 = st.text_input("🔒 Повторите пароль *", type="password")
                assistant_name = st.text_input("🤖 Имя помощника", value="Помощник")
                user_nickname = st.text_input("👨‍🏫 Как к вам обращаться?", 
                                            placeholder="Например: Мария Ивановна")
                
                if st.form_submit_button("✅ Зарегистрироваться", use_container_width=True):
                    if not all([username, password]):
                        st.error("❌ Заполните обязательные поля")
                    elif password != password2:
                        st.error("❌ Пароли не совпадают")
                    elif len(password) < 4:
                        st.error("❌ Пароль должен быть не менее 4 символов")
                    else:
                        user_id = db.register_user(username, password, 
                                                  assistant_name, user_nickname)
                        if user_id:
                            st.success("✅ Регистрация успешна! Теперь войдите в систему.")
                        else:
                            st.error("❌ Этот логин уже занят")

def main_page():
    """Главная страница с меню"""
    user = st.session_state.user
    
    # Боковое меню
    with st.sidebar:
        st.markdown(f"### 👋 {user.get('user_nickname', user['username'])}")
        st.markdown(f"🤖 Помощник: **{user.get('assistant_name', 'Помощник')}**")
        st.markdown("---")
        
        # Навигация
        menu = st.radio(
            "📋 Меню:",
            ["💬 Чат с помощником", "📅 Календарь", "📝 Заметки", 
             "📚 Работа с КТП", "👥 Список класса", "⚙️ Настройки"]
        )
        
        # Добавьте эту строку после st.radio:
        if st.button("📊 График СОР/СОЧ", use_container_width=True):
            st.switch_page("pages/sor_soch.py")
        
        st.markdown("---")
        
        # Информация о системе
        st.info(f"""
        🎤 Голосовой ввод: {'✅ Доступен' if VOICE_AVAILABLE else '❌ Недоступен'}
        🔊 Озвучивание: {'✅ Доступно' if TTS_AVAILABLE else '❌ Недоступно'}
        """)
        
        if st.button("🚪 Выйти", use_container_width=True):
            st.session_state.user = None
            st.session_state.current_page = 'login'
            st.rerun()
    
    # Основной контент
    if menu == "💬 Чат с помощником":
        chat_page()
    elif menu == "📅 Календарь":
        calendar_page()
    elif menu == "📝 Заметки":
        notes_page()
    elif menu == "📚 Работа с КТП":
        # Используем сохраненную вкладку
        if 'ktp_tab' in st.session_state:
            ktp_page(st.session_state.ktp_tab)
        else:
            ktp_page(0)
    elif menu == "👥 Список класса":
        class_list_page()
    elif menu == "⚙️ Настройки":
        settings_page()

def chat_page():
    """Страница чата"""
    user = st.session_state.user
    assistant_name = user.get('assistant_name', 'Помощник')
    
    st.markdown(f'<div class="main-header"><h1>💬 Чат с {assistant_name}</h1></div>', 
                unsafe_allow_html=True)
    
    # Выбор режима
    col1, col2 = st.columns(2)
    with col1:
        if st.button("⌨️ Текстовый режим", use_container_width=True, 
                    type="primary" if st.session_state.chat_mode == 'text' else "secondary"):
            st.session_state.chat_mode = 'text'
            st.rerun()
    with col2:
        if st.button("🎤 Голосовой режим", use_container_width=True,
                    type="primary" if st.session_state.chat_mode == 'voice' else "secondary"):
            st.session_state.chat_mode = 'voice'
            st.rerun()
    
    # Отображение истории чата
    chat_container = st.container()
    with chat_container:
        for msg in st.session_state.chat_history:
            st.markdown(f"""
            <div class="chat-message user-message">
                <strong>Вы:</strong> {msg['message']}
            </div>
            <div class="chat-message bot-message">
                <strong>{assistant_name}:</strong> {msg['response']}
            </div>
            """, unsafe_allow_html=True)
    
    # Ввод сообщения
    user_input = None
    
    if st.session_state.chat_mode == 'voice':
        st.markdown("### 🎤 Голосовой ввод")
        
        # Информация о выбранном микрофоне
        if 'selected_microphone' in st.session_state and st.session_state.selected_microphone is not None:
            st.info(f"Используется микрофон: {st.session_state.selected_microphone}")
        
        col1, col2, col3 = st.columns([2, 2, 2])
        with col2:
            if st.button("🎤 Начать запись", use_container_width=True):
                user_input = voice_to_text(
                    st.session_state.get('selected_microphone', None)
                )
                if user_input:
                    st.success(f"Распознано: {user_input}")
    else:
        user_input = st.chat_input("Введите ваш вопрос...")
    
    if user_input:
        response, navigation = process_teacher_query(user_input, user)
        
        # Сохранение в историю
        db.save_chat_message(user['id'], user_input, response, st.session_state.chat_mode)
        st.session_state.chat_history.append({
            'message': user_input,
            'response': response
        })
        
        # Озвучивание
        if st.session_state.chat_mode == 'voice':
            text_to_speech(response)
        
        # НАВИГАЦИЯ ЧЕРЕЗ КНОПКИ
        if navigation and navigation.startswith('button_'):
            if navigation == 'button_ktp_create':
                if st.button("📚 Перейти к созданию КТП", use_container_width=True):
                    st.session_state.current_page = 'ktp'
                    st.session_state.ktp_tab = 0  # Вкладка "Создать шаблон"
                    st.rerun()
            
            elif navigation == 'button_ktp_fill':
                if st.button("📅 Перейти к заполнению дат", use_container_width=True):
                    st.session_state.current_page = 'ktp'
                    st.session_state.ktp_tab = 1  # Вкладка "Заполнить даты"
                    st.rerun()
            
            elif navigation == 'button_ktp_sor':
                if st.button("📊 Перейти к графику СОР/СОЧ", use_container_width=True):
                    st.session_state.current_page = 'ktp'
                    st.session_state.ktp_tab = 2  # Вкладка "График СОР/СОЧ"
                    st.rerun()
            
            elif navigation == 'button_class_list':
                if st.button("👥 Перейти к списку класса", use_container_width=True):
                    st.session_state.current_page = 'class_list'
                    st.rerun()
            
            elif navigation == 'button_notes':
                if st.button("📝 Перейти к заметкам", use_container_width=True):
                    st.session_state.current_page = 'notes'
                    st.rerun()
            
            elif navigation == 'button_calendar':
                if st.button("📅 Перейти к календарю", use_container_width=True):
                    st.session_state.current_page = 'calendar'
                    st.rerun()

def process_calendar_query(query):
    """Обработка запросов про календарь (Казахстан)"""
    today = date.today()
    query_lower = query.lower()
    
    # Определяем текущую четверть
    current_quarter = calendar_helper.get_current_quarter()
    
    if current_quarter and current_quarter > 0:
        quarter_info = calendar_helper.get_quarter_info(current_quarter)
        
        if quarter_info:
            start = quarter_info['start_date']
            end = quarter_info['end_date']
            
            # Разные типы запросов
            if 'рабоч' in query_lower or 'учебн' in query_lower:
                return f"""📊 Информация о {current_quarter}-й четверти:
• Начало: {start.strftime('%d.%m.%Y')}
• Конец: {end.strftime('%d.%m.%Y')}
• Рабочих дней осталось: {quarter_info['working_days_left']}
• Всего рабочих дней: {quarter_info['working_days_count']}
• Праздников: {len(quarter_info['holidays'])}"""
            
            elif 'праздник' in query_lower or 'выходн' in query_lower:
                if quarter_info['holidays']:
                    holidays_text = "\n".join([
                        f"• {h['date'].strftime('%d.%m.%Y')} - {h['name']}"
                        for h in quarter_info['holidays']
                    ])
                    return f"🎉 Праздники в {current_quarter}-й четверти:\n{holidays_text}"
                else:
                    return f"В {current_quarter}-й четверти нет официальных праздников."
            
            elif 'каникул' in query_lower:
                if quarter_info['holiday_start']:
                    return f"""🏖️ Каникулы после {current_quarter}-й четверти:
• Начало: {quarter_info['holiday_start'].strftime('%d.%m.%Y')}
• Конец: {quarter_info['holiday_end'].strftime('%d.%m.%Y')}
• Длительность: {(quarter_info['holiday_end'] - quarter_info['holiday_start']).days + 1} дней"""
                else:
                    return "Летние каникулы!"
            
            else:
                # Общий ответ
                return f"""📅 Информация о {current_quarter}-й четверти:
• Начало: {start.strftime('%d.%m.%Y')}
• Конец: {end.strftime('%d.%m.%Y')}
• Осталось дней: {quarter_info['days_left']}
• Осталось недель: {quarter_info['weeks_left']}
• Осталось рабочих дней: {quarter_info['working_days_left']}"""
    
    elif current_quarter == 0:
        return """☀️ Сейчас летние каникулы!
Новый учебный год начнется 1 сентября."""
    
    return "Не удалось определить информацию о четверти"

def process_calendar_query(query):
    """Обработка запросов про календарь"""
    today = date.today()
    query_lower = query.lower()
    
    # Определяем текущую четверть
    current_quarter = calendar_helper.get_current_quarter()
    
    # Если спрашивают про конкретную четверть
    for q in range(1, 5):
        if f"{q} четверть" in query_lower or f"{q}-я четверть" in query_lower or f"{q}я четверть" in query_lower:
            info = calendar_helper.get_quarter_info(q)
            if info:
                return format_quarter_response(info)
    
    # Если спрашивают про текущую четверть
    if current_quarter and isinstance(current_quarter, int) and current_quarter > 0:
        info = calendar_helper.get_quarter_info(current_quarter)
        if info:
            return format_quarter_response(info)
    
    # Если каникулы
    if current_quarter == 0:
        return "☀️ Сейчас летние каникулы! Новая четверть начнется 1 сентября."
    elif current_quarter and isinstance(current_quarter, str) and 'каникулы' in current_quarter:
        q_num = int(current_quarter.split('_')[1])
        next_q = q_num + 1 if q_num < 4 else None
        
        if next_q:
            info = calendar_helper.get_quarter_info(next_q)
            if info:
                return f"🏖️ Сейчас каникулы! {format_quarter_response(info)}"
    
    return "Не удалось определить информацию о четверти"

def format_quarter_response(info):
    """Форматировать ответ о четверти"""
    if info['days_until_start'] > 0:
        return f"""⏳ {info['quarter']}-я четверть еще не началась:
• Начало: {info['start_date'].strftime('%d.%m.%Y')}
• До начала: {info['days_until_start']} дней ({info['weeks_until_start']} недель)
• Всего рабочих дней: {info['working_days_count']}"""
    elif info['days_until_end'] >= 0:
        return f"""📖 {info['quarter']}-я четверть идет:
• Начало: {info['start_date'].strftime('%d.%m.%Y')}
• Конец: {info['end_date'].strftime('%d.%m.%Y')}
• До конца: {info['days_until_end']} дней ({info['weeks_until_end']} недель)
• Осталось рабочих дней: {info['working_days_left']}"""
    else:
        return f"""✅ {info['quarter']}-я четверть завершена:
• Закончилась: {info['end_date'].strftime('%d.%m.%Y')}"""

def calendar_page():
    st.markdown('<div class="main-header"><h1>📅 Учебный календарь 2026-2027</h1></div>', 
                unsafe_allow_html=True)
    
    # Информация
    col1, col2, col3 = st.columns(3)
    with col1:
        st.metric("📚 Учебный год", "2026-2027")
    with col2:
        current_quarter = calendar_helper.get_current_quarter()
        if current_quarter == 0:
            quarter_name = "☀️ Летние каникулы"
        elif current_quarter and isinstance(current_quarter, str) and 'каникулы' in current_quarter:
            q_num = current_quarter.split('_')[1]
            quarter_name = f"🏖️ Каникулы (после {q_num}-й четверти)"
        elif current_quarter:
            quarter_name = f"📖 {current_quarter}-я четверть"
        else:
            quarter_name = "🏖️ Каникулы"
        st.metric("Текущий период", quarter_name)
    with col3:
        st.metric("📅 Сегодня", date.today().strftime('%d.%m.%Y'))
    
    st.markdown("---")
    
    tab1, tab2, tab3, tab4 = st.tabs([
        "⏱️ Обратный отсчет",
        "📅 Даты по дням",
        "🎉 Праздники",
        "📋 Все четверти"
    ])
    
    with tab1:
        st.subheader("⏱️ Обратный отсчет до четвертей 2026-2027")
        
        # Выбор: все четверти или конкретная
        view_mode = st.radio(
            "Режим просмотра:",
            ["Все четверти", "Конкретная четверть"],
            horizontal=True
        )
        
        if view_mode == "Конкретная четверть":
            quarter = st.selectbox(
                "Выберите четверть:",
                [1, 2, 3, 4],
                format_func=lambda x: f"{x}-я четверть"
            )
            
            info = calendar_helper.get_quarter_info(quarter)
            
            if info:
                if info['days_until_start'] > 0:
                    st.markdown(f"### ⏳ До начала {quarter}-й четверти")
                    col1, col2, col3, col4 = st.columns(4)
                    with col1:
                        st.metric("Дней", info['days_until_start'])
                    with col2:
                        st.metric("Недель", info['weeks_until_start'])
                    with col3:
                        st.metric("Рабочих дней", info['working_days_count'])
                    with col4:
                        st.metric("Начало", info['start_date'].strftime('%d.%m.%Y'))
                
                elif info['days_until_end'] >= 0:
                    st.markdown(f"### 📖 {quarter}-я четверть идет")
                    col1, col2, col3, col4 = st.columns(4)
                    with col1:
                        st.metric("До конца (дней)", info['days_until_end'])
                    with col2:
                        st.metric("До конца (недель)", info['weeks_until_end'])
                    with col3:
                        st.metric("Рабочих дней осталось", info['working_days_left'])
                    with col4:
                        st.metric("Конец", info['end_date'].strftime('%d.%m.%Y'))
                    
                    total_days = info['total_days']
                    days_passed = total_days - info['days_until_end']
                    progress = days_passed / total_days if total_days > 0 else 0
                    st.progress(min(progress, 1.0), text=f"Пройдено {days_passed} из {total_days} дней")
        
        else:
            # Показываем все четверти
            all_quarters = calendar_helper.get_all_quarters_info()
            
            for info in all_quarters:
                q = info['quarter']
                
                if info['days_until_start'] > 0:
                    # Еще не началась
                    st.markdown(f"### ⏳ {q}-я четверть (начнется через {info['days_until_start']} дней)")
                    col1, col2, col3 = st.columns(3)
                    with col1:
                        st.metric("Начало", info['start_date'].strftime('%d.%m.%Y'))
                    with col2:
                        st.metric("Конец", info['end_date'].strftime('%d.%m.%Y'))
                    with col3:
                        st.metric("Рабочих дней", info['working_days_count'])
                
                elif info['days_until_end'] >= 0:
                    # Идет
                    st.markdown(f"### 📖 {q}-я четверть (осталось {info['days_until_end']} дней)")
                    col1, col2, col3 = st.columns(3)
                    with col1:
                        st.metric("Начало", info['start_date'].strftime('%d.%m.%Y'))
                    with col2:
                        st.metric("Конец", info['end_date'].strftime('%d.%m.%Y'))
                    with col3:
                        st.metric("Осталось раб. дней", info['working_days_left'])
                    
                    total_days = info['total_days']
                    days_passed = total_days - info['days_until_end']
                    progress = days_passed / total_days if total_days > 0 else 0
                    st.progress(min(progress, 1.0), text=f"{q}-я четверть: {days_passed}/{total_days} дней")
                
                st.markdown("---")
    
    with tab2:
        st.subheader("Даты уроков по дням недели")
        
        quarter = st.selectbox(
            "Четверть:",
            [1, 2, 3, 4],
            format_func=lambda x: f"{x}-я четверть",
            key="dates_quarter"
        )
        
        start, end = calendar_helper.get_quarter_dates(quarter)
        
        if start and end:
            days_names = ['ПН', 'ВТ', 'СР', 'ЧТ', 'ПТ', 'СБ']
            selected_days = []
            
            cols = st.columns(6)
            for i, (col, day) in enumerate(zip(cols, days_names)):
                with col:
                    if st.checkbox(day, key=f"date_day_{i}"):
                        selected_days.append(i)
            
            include_holidays = st.checkbox("Показывать праздники", value=False)
            
            if selected_days:
                dates = calendar_helper.get_dates_by_weekdays(
                    start, end, selected_days, include_holidays
                )
                
                if dates:
                    df = calendar_helper.format_dates_table(dates)
                    st.dataframe(df, use_container_width=True)
                    
                    total = sum(len(d) for d in dates.values())
                    st.success(f"Всего уроков: {total}")
    
    with tab3:
        st.subheader("График СОР и СОЧ")
        
        uploaded_file = st.file_uploader("📁 Загрузите КТП", type=['docx'], key="sor_final")
        
        if uploaded_file:
            doc = Document(BytesIO(uploaded_file.read()))
            
            results = []
            
            for table in doc.tables:
                for row_idx in range(1, len(table.rows)):
                    row = table.rows[row_idx]
                    
                    try:
                        topic = row.cells[2].text
                        dates = row.cells[5].text
                        note = row.cells[6].text if len(row.cells) > 6 else ""
                    except:
                        continue
                    
                    # Объединяем тему и примечание
                    full_text = topic + " " + note
                    
                    # ПОИСК СОР
                    sor_nums = re.findall(r'сор[^\d]*(\d+)', full_text.lower())
                    # ПОИСК СОЧ
                    soch_nums = re.findall(r'соч[^\d]*(\d+)', full_text.lower())
                    
                    if sor_nums or soch_nums:
                        # Разбираем даты по классам
                        for line in dates.split('\n'):
                            m = re.match(r'(\d+[а-яА-ЯёЁ]*)\s*-\s*([\d\.\-]+)', line.strip())
                            if m:
                                cls = m.group(1)
                                dt = m.group(2)
                                
                                # Определяем четверть
                                quarter = "Не определена"
                                try:
                                    parts = dt.split('.')
                                    if len(parts) >= 2:
                                        day = int(parts[0])
                                        month = int(parts[1])
                                        d = date(2026 if month >= 9 else 2027, month, day)
                                        
                                        if date(2026, 9, 1) <= d <= date(2026, 10, 25):
                                            quarter = "1 четверть"
                                        elif date(2026, 11, 2) <= d <= date(2026, 12, 29):
                                            quarter = "2 четверть"
                                        elif date(2027, 1, 11) <= d <= date(2027, 3, 21):
                                            quarter = "3 четверть"
                                        elif date(2027, 3, 29) <= d <= date(2027, 5, 25):
                                            quarter = "4 четверть"
                                except:
                                    pass
                                
                                for num in sor_nums:
                                    results.append({
                                        'Класс': cls,
                                        'Четверть': quarter,
                                        'Тип': 'СОР',
                                        'Номер': num,
                                        'Дата': dt if dt != '-' else '-'
                                    })
                                
                                for num in soch_nums:
                                    results.append({
                                        'Класс': cls,
                                        'Четверть': quarter,
                                        'Тип': 'СОЧ',
                                        'Номер': num,
                                        'Дата': dt if dt != '-' else '-'
                                    })
            
            if results:
                # Сортировка
                def sort_key(item):
                    cls = item['Класс']
                    dt = item['Дата']
                    try:
                        parts = dt.split('.')
                        if len(parts) >= 2:
                            d = int(parts[0])
                            m = int(parts[1])
                            y = 2026 if m >= 9 else 2027
                            return (cls, y, m, d)
                    except:
                        pass
                    return (cls, 9999, 99, 99)
                
                results.sort(key=sort_key)
                
                df = pd.DataFrame(results)
                st.success(f"✅ Найдено {len(df)} записей СОР/СОЧ")
                st.dataframe(df, use_container_width=True)
                
                # Статистика
                col1, col2, col3 = st.columns(3)
                with col1:
                    st.metric("СОР", len(df[df['Тип'] == 'СОР']))
                with col2:
                    st.metric("СОЧ", len(df[df['Тип'] == 'СОЧ']))
                with col3:
                    st.metric("Классов", len(df['Класс'].unique()))
                
                # Excel
                buffer = io.BytesIO()
                with pd.ExcelWriter(buffer, engine='openpyxl') as writer:
                    df.to_excel(writer, sheet_name='Общий график', index=False)
                    for cls in sorted(df['Класс'].unique()):
                        cls_df = df[df['Класс'] == cls]
                        sheet_name = f"Класс_{cls}"[:31]
                        cls_df.to_excel(writer, sheet_name=sheet_name, index=False)
                buffer.seek(0)
                
                st.download_button(
                    label="📥 Скачать Excel",
                    data=buffer,
                    file_name="график_СОР_СОЧ.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    use_container_width=True
                )
            else:
                st.warning("СОР/СОЧ не найдены")
    
    with tab4:
        st.subheader("Сводная таблица четвертей")
        
        all_quarters = calendar_helper.get_all_quarters_info()
        
        data = []
        for info in all_quarters:
            data.append({
                'Четверть': f"{info['quarter']}-я",
                'Начало': info['start_date'].strftime('%d.%m.%Y'),
                'Конец': info['end_date'].strftime('%d.%m.%Y'),
                'Дней': info['total_days'],
                'Рабочих дней': info['working_days_count'],
                'Праздников': len(info['holidays'])
            })
        
        if data:
            df = pd.DataFrame(data)
            st.dataframe(df, use_container_width=True)
        
        # Каникулы
        st.markdown("### 🏖️ Каникулы")
        for info in all_quarters:
            if info['holiday_start']:
                days = (info['holiday_end'] - info['holiday_start']).days + 1
                st.markdown(f"**После {info['quarter']}-й четверти:** {info['holiday_start'].strftime('%d.%m.%Y')} — {info['holiday_end'].strftime('%d.%m.%Y')} ({days} дней)")

def notes_page():
    """Страница заметок с календарем"""
    user = st.session_state.user
    
    st.markdown('<div class="main-header"><h1>📝 Заметки и календарь</h1></div>', 
                unsafe_allow_html=True)
    
    # Инициализация календаря заметок
    notes_calendar = NotesCalendar()
    
    # CSS для красивого календаря
    st.markdown("""
    <style>
        .calendar-container {
            background: white;
            border-radius: 15px;
            padding: 20px;
            box-shadow: 0 4px 6px rgba(0,0,0,0.1);
        }
        
        .calendar-header {
            text-align: center;
            font-size: 24px;
            font-weight: bold;
            margin-bottom: 20px;
            color: #333;
        }
        
        .calendar-day {
            min-height: 80px;
            border: 1px solid #e0e0e0;
            border-radius: 8px;
            padding: 8px;
            margin: 2px;
            cursor: pointer;
            transition: all 0.3s;
        }
        
        .calendar-day:hover {
            box-shadow: 0 2px 8px rgba(0,0,0,0.15);
            transform: translateY(-2px);
        }
        
        .calendar-day-today {
            border: 2px solid #667eea;
            background: #f0f2ff;
        }
        
        .calendar-day-weekend {
            background: #f8f9fa;
        }
        
        .calendar-day-holiday {
            background: #FFF1B8;
        }
        
        .calendar-day-has-notes {
            border-left: 4px solid #667eea;
        }
        
        .note-chip {
            display: inline-block;
            padding: 4px 8px;
            border-radius: 12px;
            font-size: 11px;
            margin: 2px;
            cursor: pointer;
        }
        
        .note-chip-high {
            background: #FFD6D6;
            color: #c0392b;
        }
        
        .note-chip-normal {
            background: #FFE4C2;
            color: #e67e22;
        }
        
        .note-chip-low {
            background: #D9F2DF;
            color: #27ae60;
        }
        
        .note-chip-holiday {
            background: #FFF1B8;
            color: #f39c12;
        }
        
        .week-block {
            background: #f8f9fa;
            border-radius: 15px;
            padding: 15px;
            margin: 10px 0;
        }
        
        .week-day {
            background: white;
            border-radius: 10px;
            padding: 10px;
            margin: 5px 0;
            border: 1px solid #e0e0e0;
        }
        
        .week-day-today {
            border: 2px solid #667eea;
            background: #f0f2ff;
        }
        
        .week-day-weekend {
            background: #f8f9fa;
        }
        
        .note-item {
            padding: 8px;
            border-radius: 8px;
            margin: 5px 0;
            font-size: 14px;
        }
        
        .note-item-high {
            background: #FFD6D6;
            border-left: 3px solid #e74c3c;
        }
        
        .note-item-normal {
            background: #FFE4C2;
            border-left: 3px solid #f39c12;
        }
        
        .note-item-low {
            background: #D9F2DF;
            border-left: 3px solid #27ae60;
        }
        
        @media (max-width: 768px) {
            .calendar-day {
                min-height: 50px;
                font-size: 12px;
                padding: 4px;
            }
        }
    </style>
    """, unsafe_allow_html=True)
    
    # Вкладки
    tab1, tab2, tab3, tab4 = st.tabs([
        "📅 Календарь", 
        "📝 Все заметки", 
        "➕ Добавить", 
        "🔔 Напоминания"
    ])
    
    with tab1:
        st.subheader("📅 Календарь заметок")
        
        # Навигация по месяцам
        col1, col2, col3 = st.columns([1, 3, 1])
        
        with col1:
            if 'calendar_month' not in st.session_state:
                st.session_state.calendar_month = date.today().month
            if 'calendar_year' not in st.session_state:
                st.session_state.calendar_year = date.today().year
            
            if st.button("⬅️ Пред.", use_container_width=True):
                st.session_state.calendar_month -= 1
                if st.session_state.calendar_month < 1:
                    st.session_state.calendar_month = 12
                    st.session_state.calendar_year -= 1
                st.rerun()
        
        with col2:
            st.markdown(f"### {notes_calendar.month_names_ru[st.session_state.calendar_month]} {st.session_state.calendar_year}")
        
        with col3:
            if st.button("След. ➡️", use_container_width=True):
                st.session_state.calendar_month += 1
                if st.session_state.calendar_month > 12:
                    st.session_state.calendar_month = 1
                    st.session_state.calendar_year += 1
                st.rerun()
        
        # Получаем календарь на месяц
        month_days = notes_calendar.get_month_calendar(
            st.session_state.calendar_year, 
            st.session_state.calendar_month
        )
        
        # Получаем все заметки с датами
        all_dated_notes = db.get_all_notes_with_dates(user['id'])
        notes_by_date = {}
        for note in all_dated_notes:
            if note.get('reminder_date'):
                date_key = note['reminder_date'][:10]  # YYYY-MM-DD
                if date_key not in notes_by_date:
                    notes_by_date[date_key] = []
                notes_by_date[date_key].append(note)
        
        # Заголовки дней недели
        cols = st.columns(7)
        for i, day_name in enumerate(notes_calendar.day_names_ru):
            with cols[i]:
                st.markdown(f"**{day_name}**")
        
        # Отображение календаря
        for week in month_days:
            cols = st.columns(7)
            
            for i, day in enumerate(week):
                with cols[i]:
                    if day == 0:
                        st.markdown("")
                    else:
                        current_date = date(
                            st.session_state.calendar_year,
                            st.session_state.calendar_month,
                            day
                        )
                        
                        # Определяем классы для стилизации
                        is_today = notes_calendar.is_today(current_date)
                        is_weekend = notes_calendar.is_weekend(current_date)
                        date_key = current_date.strftime('%Y-%m-%d')
                        has_notes = date_key in notes_by_date
                        is_holiday = calendar_helper.is_holiday(current_date)[0]
                        
                        # Формируем HTML
                        classes = ['calendar-day']
                        if is_today:
                            classes.append('calendar-day-today')
                        if is_weekend:
                            classes.append('calendar-day-weekend')
                        if is_holiday:
                            classes.append('calendar-day-holiday')
                        if has_notes:
                            classes.append('calendar-day-has-notes')
                        
                        class_str = ' '.join(classes)
                        
                        # Информация о заметках (точками)
                        notes_html = ""
                        if has_notes:
                            # Добавляем точки вместо текста
                            dots_html = ""
                            for note in notes_by_date[date_key][:5]:
                                priority = note.get('priority', 'normal')
                                color = {
                                    'high': '#FF4444',
                                    'normal': '#FFAA00',
                                    'low': '#00AA00'
                                }.get(priority, '#999')
                                dots_html += f'<span style="display:inline-block; width:8px; height:8px; border-radius:50%; background:{color}; margin:1px;"></span>'
                            
                            if len(notes_by_date[date_key]) > 5:
                                dots_html += f'<span style="font-size:10px;">+{len(notes_by_date[date_key]) - 5}</span>'
                            
                            notes_html = dots_html
                        
                        # Праздник
                        holiday_text = ""
                        if is_holiday:
                            _, holiday_name = calendar_helper.is_holiday(current_date)
                            holiday_text = f'<span class="note-chip note-chip-holiday">🇰🇿 {holiday_name}</span>'
                        
                        st.markdown(f"""
                        <div class="{class_str}">
                            <strong>{day}</strong>
                            {holiday_text}
                            {notes_html}
                        </div>
                        """, unsafe_allow_html=True)
        
        # Легенда
        st.markdown("---")
        st.markdown("**Легенда:**")
        col1, col2, col3, col4 = st.columns(4)
        with col1:
            st.markdown("🔴 Высокий приоритет")
        with col2:
            st.markdown("🟠 Средний приоритет")
        with col3:
            st.markdown("🟢 Низкий приоритет")
        with col4:
            st.markdown("🇰🇿 Праздник")
    
    with tab2:
        st.subheader("📝 Все заметки")
        
        all_notes = db.get_user_notes(user['id'])
        
        if not all_notes:
            st.info("У вас пока нет заметок. Создайте первую!")
        else:
            # Фильтр по приоритету
            filter_priority = st.selectbox(
                "Фильтр по приоритету:",
                ['all', 'high', 'normal', 'low'],
                format_func=lambda x: {
                    'all': 'Все',
                    'high': '🔴 Высокий',
                    'normal': '🟠 Средний',
                    'low': '🟢 Низкий'
                }[x]
            )
            
            filtered_notes = all_notes if filter_priority == 'all' else \
                           [n for n in all_notes if n.get('priority') == filter_priority]
            
            for note in filtered_notes:
                priority = note.get('priority', 'normal')
                emoji = notes_calendar.get_priority_emoji(priority)
                
                with st.container():
                    st.markdown(f"""
                    <div class="note-item note-item-{priority}">
                        <strong>{emoji} {note['title']}</strong><br>
                        {note.get('content', '')}
                        {f'<br><small>🔔 Напоминание: {note["reminder_date"]}</small>' if note.get('reminder_date') else ''}
                        <br><small>📅 Создано: {note['created_at']}</small>
                    </div>
                    """, unsafe_allow_html=True)
                    
                    col1, col2, col3 = st.columns([1, 1, 3])
                    with col1:
                        if st.button("✅", key=f"done_{note['id']}", help="Отметить как выполненное"):
                            db.update_note(note['id'], is_completed=True)
                            st.rerun()
                    with col2:
                        if st.button("🗑️", key=f"del_{note['id']}", help="Удалить"):
                            db.delete_note(note['id'])
                            st.rerun()
    
    with tab3:
        st.subheader("➕ Добавить заметку")
        
        # Выбор способа ввода
        input_mode = st.radio(
            "Способ ввода:",
            ["✏️ Текст", "🎤 Голос"],
            horizontal=True
        )
        
        if input_mode == "🎤 Голос":
            if VOICE_AVAILABLE:
                st.info("Нажмите кнопку и произнесите заметку")
                
                if st.button("🎤 Начать запись", use_container_width=True):
                    recognized_text = voice_to_text()
                    if recognized_text:
                        st.session_state.voice_note_text = recognized_text
                        st.success(f"Распознано: {recognized_text}")
                    else:
                        st.error("Не удалось распознать речь")
                
                # Поле для текста (можно редактировать после распознавания)
                note_content = st.text_area(
                    "Содержание заметки:",
                    value=st.session_state.get('voice_note_text', ''),
                    height=100
                )
            else:
                st.warning("Голосовой ввод недоступен")
                note_content = st.text_area("Содержание заметки:", height=100)
        else:
            note_content = st.text_area("Содержание заметки:", height=100)
        
        with st.form("add_note_form"):
            title = st.text_input("📌 Заголовок")
            
            col1, col2 = st.columns(2)
            with col1:
                reminder_date = st.date_input("📅 Дата напоминания", value=None)
                reminder_time = st.time_input("⏰ Время", value=datetime.now().time())
            with col2:
                priority = st.selectbox(
                    "⭐ Приоритет",
                    ['high', 'normal', 'low'],
                    format_func=lambda x: {
                        'high': '🔴 Высокий',
                        'normal': '🟠 Средний',
                        'low': '🟢 Низкий'
                    }[x]
                )
            
            submitted = st.form_submit_button("💾 Сохранить", use_container_width=True)
            
            if submitted:
                if title or note_content:
                    # Формируем полную дату с временем
                    full_datetime = datetime.combine(reminder_date, reminder_time).strftime('%Y-%m-%d %H:%M:%S') if reminder_date else None
                    
                    db.add_note(
                        user['id'], 
                        title or note_content[:30], 
                        note_content, 
                        full_datetime, 
                        priority
                    )
                    st.success("✅ Заметка сохранена!")
                    st.session_state.voice_note_text = ''
                    st.rerun()
                else:
                    st.warning("Введите заголовок или содержание")
    
    with tab4:
        st.subheader("🔔 Напоминания на неделю")
        
        # Получаем заметки на ближайшие 7 дней
        upcoming_notes = db.get_upcoming_reminders(user['id'], 7)
        
        # Текущая неделя
        today = date.today()
        week_dates = notes_calendar.get_week_dates(today)
        
        # Группируем по дням
        notes_by_weekday = {}
        for note in upcoming_notes:
            if note.get('reminder_date'):
                date_key = note['reminder_date'][:10]
                if date_key not in notes_by_weekday:
                    notes_by_weekday[date_key] = []
                notes_by_weekday[date_key].append(note)
        
        # Отображение недельного блока
        col1, col2 = st.columns([2, 1])
        
        with col1:
            st.markdown("### 📅 Текущая неделя")
            
            for day_date in week_dates:
                date_key = day_date.strftime('%Y-%m-%d')
                is_today = day_date == today
                is_weekend = day_date.weekday() >= 5
                
                day_class = 'week-day'
                if is_today:
                    day_class += ' week-day-today'
                if is_weekend:
                    day_class += ' week-day-weekend'
                
                day_name = notes_calendar.day_names_full[day_date.weekday()]
                
                notes_for_day = notes_by_weekday.get(date_key, [])
                
                notes_html = ""
                if notes_for_day:
                    for note in notes_for_day:
                        priority = note.get('priority', 'normal')
                        emoji = notes_calendar.get_priority_emoji(priority)
                        notes_html += f'<div class="note-item note-item-{priority}">{emoji} {note["title"]}</div>'
                else:
                    notes_html = '<small style="color: #999;">Нет заметок</small>'
                
                st.markdown(f"""
                <div class="{day_class}">
                    <strong>{day_name} {day_date.strftime('%d.%m')}</strong>
                    {notes_html}
                </div>
                """, unsafe_allow_html=True)
        
        with col2:
            st.markdown("### 📋 Сводка")
            
            total_notes = len(upcoming_notes)
            high_priority = len([n for n in upcoming_notes if n.get('priority') == 'high'])
            normal_priority = len([n for n in upcoming_notes if n.get('priority') == 'normal'])
            low_priority = len([n for n in upcoming_notes if n.get('priority') == 'low'])
            
            st.metric("Всего заметок", total_notes)
            st.markdown(f"🔴 Высокий приоритет: {high_priority}")
            st.markdown(f"🟠 Средний приоритет: {normal_priority}")
            st.markdown(f"🟢 Низкий приоритет: {low_priority}")
            
            # Ближайшая заметка
            if upcoming_notes:
                next_note = upcoming_notes[0]
                st.markdown("---")
                st.markdown("### ⏰ Ближайшая:")
                st.markdown(f"**{next_note['title']}**")
                st.markdown(f"📅 {next_note['reminder_date']}")

def ktp_page():
    st.markdown('<div class="main-header"><h1>📚 Работа с КТП</h1></div>', unsafe_allow_html=True)
    
    ktp_helper.set_calendar_helper(calendar_helper)
    
    # Определяем активную вкладку
    if 'ktp_active_tab' not in st.session_state:
        st.session_state.ktp_active_tab = 0
    
    tab1, tab2, tab3, tab4 = st.tabs([
        "📝 Создать шаблон",
        "📅 Заполнить даты",
        "📊 График СОР/СОЧ",
        "📖 Инструкция"
    ])
    
    # ===== ВКЛАДКА 1: СОЗДАНИЕ ШАБЛОНА =====
    with tab1:
        st.subheader("Создание шаблона КТП")
        
        st.info("""
        Создается Word-документ с альбомной ориентацией и разбивкой по четвертям.
        """)
        
        col1, col2 = st.columns(2)
        with col1:
            subject = st.text_input("📖 Предмет", placeholder="Например: Математика")
            class_name = st.text_input("🏫 Класс", placeholder="Например: 7")
            total_hours = st.number_input("⏱️ Всего часов в году", min_value=1, max_value=500, value=136)
        with col2:
            hours_per_week = st.number_input("📅 Часов в неделю", min_value=1, max_value=10, value=4)
        
        if st.button("📄 Создать Word-файл КТП", use_container_width=True):
            if subject and class_name:
                try:
                    doc, table = ktp_helper.create_ktp_template(
                        subject, class_name, total_hours, hours_per_week
                    )
                    
                    buffer = io.BytesIO()
                    doc.save(buffer)
                    buffer.seek(0)
                    
                    st.success("✅ Шаблон КТП создан!")
                    
                    st.download_button(
                        label="📥 Скачать шаблон КТП",
                        data=buffer,
                        file_name=f"КТП_{subject}_{class_name}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                except Exception as e:
                    st.error(f"Ошибка: {e}")
            else:
                st.warning("Заполните все поля")
    
    # ===== ВКЛАДКА 2: ЗАПОЛНЕНИЕ ДАТ =====
    with tab2:
        st.subheader("Заполнение дат")
        
        st.info("""
        **Как это работает:**
        1. Загрузите заполненный Word-файл КТП
        2. Укажите классы (литеры)
        3. Настройте расписание уроков
        4. Даты заполнятся автоматически
        """)
        
        uploaded_file = st.file_uploader("📁 Загрузите Word-файл КТП", type=['docx', 'doc'], key="ktp_fill_upload")
        
        if uploaded_file:
            try:
                file_extension = uploaded_file.name.split('.')[-1]
                file_bytes = uploaded_file.read()
                doc = ktp_helper.read_document(file_bytes, file_extension)
                
                is_valid, message = validate_ktp_structure(doc)
                
                if not is_valid:
                    st.error(f"❌ {message}")
                else:
                    st.success(f"✅ {message}")
                    
                    # Настройка классов
                    st.markdown("### 🏫 Настройка классов")
                    
                    col1, col2 = st.columns(2)
                    with col1:
                        base_class = st.text_input("Номер класса (параллель)", value="5")
                    with col2:
                        letters_input = st.text_input(
                            "Введите литеры через запятую",
                            value="а, б",
                            placeholder="Например: а, б, в"
                        )
                    
                    classes_list = []
                    for letter in letters_input.split(','):
                        letter = letter.strip()
                        if letter:
                            classes_list.append(f"{base_class}{letter}")
                    
                    st.write("Будут заполнены даты для:")
                    for cls in classes_list:
                        st.markdown(f"- **{cls}**")
                    
                    # Расписание
                    st.markdown("### 📅 Расписание уроков")
                    
                    hours_per_week = st.number_input(
                        "Сколько уроков в неделю?",
                        min_value=1, max_value=5, value=2
                    )
                    
                    days_names = ['ПН', 'ВТ', 'СР', 'ЧТ', 'ПТ']
                    day_numbers = [0, 1, 2, 3, 4]
                    
                    class_configs = {}
                    
                    for cls in classes_list:
                        st.markdown(f"#### {cls}")
                        
                        lessons = []
                        
                        for lesson_num in range(1, hours_per_week + 1):
                            st.markdown(f"**{lesson_num} урок:**")
                            
                            cols = st.columns(5)
                            selected_days = []
                            
                            for i, (col, day_name) in enumerate(zip(cols, days_names)):
                                with col:
                                    if st.checkbox(day_name, value=False, key=f"{cls}_{lesson_num}_{i}"):
                                        selected_days.append(day_numbers[i])
                            
                            if selected_days:
                                lessons.append({'weekday': selected_days[0], 'label': f'{lesson_num} урок'})
                            else:
                                default_day = day_numbers[(lesson_num - 1) % 5]
                                lessons.append({'weekday': default_day, 'label': f'{lesson_num} урок'})
                                st.caption(f"По умолчанию: {days_names[default_day]}")
                        
                        class_configs[cls] = {'lessons': lessons}
                    
                    # Кнопка
                    if st.button("📅 Заполнить даты", use_container_width=True, type="primary"):
                        try:
                            with st.spinner("Заполняю даты..."):
                                result = ktp_helper.fill_dates_for_classes(doc, class_configs)
                                
                                if len(result) == 5:
                                    doc, lessons_count, compact_warnings, extra_warnings, debug_info = result
                                else:
                                    doc, lessons_count, compact_warnings, extra_warnings = result
                                    debug_info = None
                                
                                st.success(f"✅ Запланировано тем: {lessons_count}")
                                
                                # Предупреждения
                                if compact_warnings:
                                    st.markdown("""
                                    <div style="background-color: #FFE4E1; padding: 15px; border-radius: 10px; border-left: 5px solid #E74C3C;">
                                        <h4 style="color: #E74C3C; margin: 0;">⚠️ Требуется уплотнение</h4>
                                    </div>
                                    """, unsafe_allow_html=True)
                                    
                                    quarter_names = {1: 'I', 2: 'II', 3: 'III', 4: 'IV'}
                                    for cls, quarters_data in compact_warnings.items():
                                        if isinstance(quarters_data, dict):
                                            for q, count in quarters_data.items():
                                                q_name = quarter_names.get(q, str(q))
                                                st.warning(f"**{cls}** — {q_name} четверть: уплотнить {count} час(ов)")
                                
                                if extra_warnings:
                                    st.markdown("""
                                    <div style="background-color: #E0F7FA; padding: 15px; border-radius: 10px; border-left: 5px solid #00BCD4;">
                                        <h4 style="color: #00838F; margin: 0;">📋 Свободные даты</h4>
                                    </div>
                                    """, unsafe_allow_html=True)
                                    
                                    for cls, quarters_data in extra_warnings.items():
                                        if isinstance(quarters_data, dict):
                                            for q, data in quarters_data.items():
                                                q_name = quarter_names.get(q, str(q))
                                                if isinstance(data, dict):
                                                    count = data.get('count', 0)
                                                    st.info(f"**{cls}** — {q_name} четверть: {count} свободных дат")
                                
                                # Скачивание
                                buffer = io.BytesIO()
                                doc.save(buffer)
                                buffer.seek(0)
                                
                                original_name = uploaded_file.name.rsplit('.', 1)[0]
                                new_filename = f"{original_name} с датами.docx"
                                
                                st.download_button(
                                    label=f"📥 Скачать {new_filename}",
                                    data=buffer,
                                    file_name=new_filename,
                                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                                )
                        except Exception as e:
                            st.error(f"Ошибка: {e}")
                            import traceback
                            st.code(traceback.format_exc())
            
            except Exception as e:
                st.error(f"Ошибка чтения файла: {e}")
    
    # ===== ВКЛАДКА 3: ГРАФИК СОР/СОЧ =====
    with tab3:
        st.subheader("График СОР и СОЧ")
        
        uploaded_file = st.file_uploader("📁 Загрузите КТП", type=['docx'], key="sor_soch_upload")
        
        if uploaded_file:
            doc = Document(BytesIO(uploaded_file.read()))
            
            results = []
            
            for table in doc.tables:
                for row_idx in range(1, len(table.rows)):
                    row = table.rows[row_idx]
                    
                    try:
                        topic = row.cells[2].text.strip()
                        dates = row.cells[5].text.strip()
                        note = row.cells[6].text.strip() if len(row.cells) > 6 else ""
                    except:
                        continue
                    
                    full_text = topic + " " + note
                    text_lower = full_text.lower()
                    
                    sor_nums = re.findall(r'сор[^\d]*(\d+)', text_lower)
                    soch_nums = re.findall(r'соч[^\d]*(\d+)', text_lower)
                    
                    if sor_nums or soch_nums:
                        for line in dates.split('\n'):
                            m = re.match(r'(\d+[а-яА-ЯёЁ]*)\s*-\s*([\d\.\-]+)', line.strip())
                            if m:
                                cls = m.group(1).strip()
                                dt = m.group(2).strip()
                                
                                quarter = "Не определена"
                                if dt and dt != '-':
                                    try:
                                        parts = dt.split('.')
                                        if len(parts) >= 2:
                                            day = int(parts[0])
                                            month = int(parts[1])
                                            d = date(2026 if month >= 9 else 2027, month, day)
                                            
                                            if date(2026, 9, 1) <= d <= date(2026, 10, 25):
                                                quarter = "1 четверть"
                                            elif date(2026, 11, 2) <= d <= date(2026, 12, 29):
                                                quarter = "2 четверть"
                                            elif date(2027, 1, 11) <= d <= date(2027, 3, 21):
                                                quarter = "3 четверть"
                                            elif date(2027, 3, 29) <= d <= date(2027, 5, 25):
                                                quarter = "4 четверть"
                                    except:
                                        pass
                                
                                for num in sor_nums:
                                    results.append({
                                        'Класс': cls,
                                        'Четверть': quarter,
                                        'Тип': 'СОР',
                                        'Номер': num,
                                        'Дата': dt if dt else '-'
                                    })
                                
                                for num in soch_nums:
                                    results.append({
                                        'Класс': cls,
                                        'Четверть': quarter,
                                        'Тип': 'СОЧ',
                                        'Номер': num,
                                        'Дата': dt if dt else '-'
                                    })
            
            if results:
                def sort_key(item):
                    cls = item['Класс']
                    dt = item['Дата']
                    try:
                        parts = dt.split('.')
                        if len(parts) >= 2:
                            d = int(parts[0])
                            m = int(parts[1])
                            y = 2026 if m >= 9 else 2027
                            return (cls, y, m, d)
                    except:
                        pass
                    return (cls, 9999, 99, 99)
                
                results.sort(key=sort_key)
                df = pd.DataFrame(results)
                
                st.success(f"✅ Найдено {len(df)} записей")
                st.dataframe(df, use_container_width=True)
                
                # Excel
                buffer = io.BytesIO()
                with pd.ExcelWriter(buffer, engine='openpyxl') as writer:
                    df.to_excel(writer, sheet_name='Общий график', index=False)
                    for cls in sorted(df['Класс'].unique()):
                        cls_df = df[df['Класс'] == cls]
                        sheet_name = f"Класс_{cls}"[:31]
                        cls_df.to_excel(writer, sheet_name=sheet_name, index=False)
                buffer.seek(0)
                
                st.download_button(
                    label="📥 Скачать Excel",
                    data=buffer,
                    file_name="график_СОР_СОЧ.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    use_container_width=True
                )
            else:
                st.warning("СОР/СОЧ не найдены")
    
    # ===== ВКЛАДКА 4: ИНСТРУКЦИЯ =====
    with tab4:
        st.subheader("Инструкция")
        st.markdown("""
        ### 📝 Шаг 1: Создание шаблона
        - Укажите предмет, класс, часы
        - Скачайте шаблон Word
        
        ### ✏️ Шаг 2: Заполнение
        Заполните: №, Раздел, Темы, Цели обучения, Кол-во часов
        
        ### 📅 Шаг 3: Заполнение дат
        1. Загрузите файл
        2. Укажите классы
        3. Настройте расписание
        4. Даты заполнятся автоматически
        
        ### 📊 Шаг 4: СОР/СОЧ
        Загрузите КТП с датами — график создастся автоматически
        """)

def class_list_page():
    """Страница работы со списком класса"""
    st.markdown('<div class="main-header"><h1>👥 Список класса</h1></div>', 
                unsafe_allow_html=True)
    
    tab1, tab2 = st.tabs(["📤 Загрузить список", "🔍 Поиск ученика"])
    
    with tab1:
        st.subheader("Загрузка списка класса из Excel")
        st.info("""
        Файл Excel должен содержать колонки:
        • Фамилия (обязательно)
        • Имя (обязательно)
        • Отчество (обязательно)
        • Адрес
        • Телефон
        • Дата рождения
        • и другие данные учеников
        """)
        
        class_name = st.text_input("Название класса", placeholder="Например: 9А")
        uploaded_file = st.file_uploader("Выберите Excel-файл", type=['xlsx', 'xls'])
        
        if uploaded_file and class_name:
            try:
                df = excel_helper.load_class_list(uploaded_file)
                st.success(f"✅ Загружено {len(df)} учеников")
                
                # Сохранение в базу
                db.save_class_list(st.session_state.user['id'], class_name, 
                                 df.to_dict('records'))
                
                # Просмотр данных
                st.subheader("📋 Превью списка")
                st.dataframe(df.head(10), use_container_width=True)
            except Exception as e:
                st.error(f"❌ Ошибка: {str(e)}")
    
    with tab2:
        st.subheader("Поиск информации об ученике")
        
        if excel_helper.class_data is None:
            # Попытка загрузить из базы
            saved_data = db.get_class_list(st.session_state.user['id'], 
                                          st.text_input("Класс для поиска", key="search_class"))
            if saved_data:
                df = pd.DataFrame(json.loads(saved_data['student_data']))
                excel_helper.class_data = df
        
        if excel_helper.class_data is not None:
            lastname = st.text_input("🔍 Фамилия ученика")
            
            if lastname:
                result = excel_helper.search_student(lastname)
                
                if result is None:
                    st.warning("Ученик не найден")
                elif isinstance(result, dict) and result.get('multiple'):
                    st.warning(f"Найдено несколько учеников с фамилией '{lastname}'")
                    st.write("Уточните имя:")
                    
                    for student in result['students']:
                        if st.button(f"{student['Фамилия']} {student['Имя']} {student.get('Отчество', '')}"):
                            info = excel_helper.get_student_info(lastname, student['Имя'])
                            if info:
                                st.markdown(info)
                else:
                    # Показываем всю информацию
                    st.subheader(f"📋 Информация об ученике")
                    for key, value in result.items():
                        if pd.notna(value):
                            st.write(f"**{key}:** {value}")
        else:
            st.info("Сначала загрузите список класса")

def settings_page():
    """Страница настроек"""
    st.markdown('<div class="main-header"><h1>⚙️ Настройки</h1></div>', 
                unsafe_allow_html=True)
    
    user = st.session_state.user
    
    st.subheader("👤 Личные данные")
    new_nickname = st.text_input("Как к вам обращаться", value=user.get('user_nickname', ''))
    new_assistant_name = st.text_input("Имя помощника", value=user.get('assistant_name', 'Помощник'))
    
    # Настройки микрофона
    st.subheader("🎤 Настройки микрофона")
    
    if VOICE_AVAILABLE:
        microphones = list_microphones()
        
        if microphones:
            st.success(f"✅ Найдено микрофонов: {len(microphones)}")
            
            # Сохраняем выбор в сессию
            if 'selected_microphone' not in st.session_state:
                st.session_state.selected_microphone = None
            
            # Создаем список микрофонов
            mic_options = ["Автоматически"] + [f"Микрофон {i}: {name}" for i, name in enumerate(microphones)]
            
            selected = st.selectbox(
                "Выберите микрофон:",
                range(len(mic_options)),
                format_func=lambda x: mic_options[x]
            )
            
            if selected == 0:
                st.session_state.selected_microphone = None
            else:
                st.session_state.selected_microphone = selected - 1  # Минус 1, т.к. первый пункт "Автоматически"
            
            # Кнопка тестирования
            if st.button("🎤 Тестировать микрофон", use_container_width=True):
                text = voice_to_text(st.session_state.selected_microphone)
                if text:
                    st.success(f"✅ Распознано: {text}")
                else:
                    st.warning("Не удалось распознать речь. Попробуйте другой микрофон.")
        else:
            st.error("❌ Микрофоны не найдены")
            st.info("""
            Проверьте:
            1. Подключен ли микрофон
            2. Разрешен ли доступ к микрофону в браузере
            3. Работает ли микрофон в других программах
            """)
    else:
        st.warning("⚠️ Модуль распознавания речи не установлен")
        st.code("pip install SpeechRecognition pyaudio")
    
    st.markdown("---")
    
    if st.button("💾 Сохранить настройки", use_container_width=True):
        conn = db.get_connection()
        conn.execute('''
            UPDATE users SET user_nickname = ?, assistant_name = ?
            WHERE id = ?
        ''', (new_nickname, new_assistant_name, user['id']))
        conn.commit()
        conn.close()
        
        st.session_state.user['user_nickname'] = new_nickname
        st.session_state.user['assistant_name'] = new_assistant_name
        st.success("✅ Настройки сохранены!")

# ГЛАВНЫЙ РОУТЕР
def main():
    # Гарантируем инициализацию
    if 'user' not in st.session_state:
        st.session_state.user = None
    
    if st.session_state.user is None:
        login_page()
    else:
        main_page()

if __name__ == "__main__":
    main()