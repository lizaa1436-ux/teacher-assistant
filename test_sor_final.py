from docx import Document
from io import BytesIO
import re
from datetime import date
from docx.oxml.ns import qn
import os

def parse_dates(date_string):
    results = []
    if not date_string or date_string.strip() == '-':
        return results
    normalized = date_string.replace(';', '\n')
    pattern = r'(\d+[а-яА-ЯёЁ]*)\s*[-–—]\s*([\d\.\-]+)'
    matches = re.findall(pattern, normalized)
    for match in matches:
        class_name = match[0].strip()
        date_val = match[1].strip()
        results.append((class_name, date_val if date_val != '-' else '-'))
    return results

def find_sor_soch(text):
    matches = []
    seen = set()
    if not text:
        return matches
    text_lower = text.lower()
    patterns = [
        (r'соч\s*№\s*(\d+)', 'СОЧ'),
        (r'соч\s+(\d+)', 'СОЧ'),
        (r'соч(\d+)', 'СОЧ'),
        (r'сор\s*№\s*(\d+)', 'СОР'),
        (r'сор\s+(\d+)', 'СОР'),
        (r'сор(\d+)', 'СОР'),
    ]
    for pattern, wtype in patterns:
        for m in re.finditer(pattern, text_lower):
            num = m.group(1)
            key = (wtype, num)
            if key not in seen:
                seen.add(key)
                matches.append({'type': wtype, 'number': num})
    return matches

def check_h_merge(row):
    try:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.find(qn('w:tcPr'))
            if tcPr is not None:
                for child in tcPr:
                    tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
                    if tag == 'gridSpan':
                        return True
        return False
    except:
        return False

# Выбор файла
file_path = input("Путь к файлу КТП: ").strip().strip('"')

if not os.path.exists(file_path):
    print(f"Файл не найден: {file_path}")
    exit()

print(f"\n✅ Файл: {file_path}")

doc = Document(file_path)
print(f"Таблиц: {len(doc.tables)}")

found_assessments = []

for table in doc.tables:
    print(f"\nТаблица: {len(table.rows)} строк, {len(table.columns)} столбцов")
    
    # Заголовки
    print("Заголовки:")
    for i, cell in enumerate(table.rows[0].cells):
        print(f"  [{i}] {cell.text.strip()}")
    
    # Обработка строк
    for row_idx in range(1, len(table.rows)):
        row = table.rows[row_idx]
        
        # Проверяем горизонтальное объединение
        has_h_merge = check_h_merge(row)
        
        if has_h_merge:
            print(f"\nСтрока {row_idx}: ПРОПУЩЕНА (горизонтальное объединение)")
            continue
        
        try:
            topic = row.cells[2].text.strip() if len(row.cells) > 2 else ""
            dates = row.cells[5].text.strip() if len(row.cells) > 5 else ""
            note = row.cells[6].text.strip() if len(row.cells) > 6 else ""
        except Exception as e:
            print(f"\nСтрока {row_idx}: ОШИБКА {e}")
            continue
        
        # Ищем СОР/СОЧ
        matches = find_sor_soch(topic + " " + note)
        
        if matches:
            print(f"\n✅ Строка {row_idx}: НАЙДЕНО!")
            print(f"  Тема: '{topic[:80]}'")
            print(f"  Даты: '{dates}'")
            print(f"  Совпадения: {matches}")
            
            class_dates = parse_dates(dates)
            print(f"  Классы и даты: {class_dates}")
            
            for m in matches:
                for cls, dt in class_dates:
                    found_assessments.append({
                        'Класс': cls,
                        'Тип': m['type'],
                        'Номер': m['number'],
                        'Дата': dt
                    })
        else:
            # Показываем только первые 10 строк без СОР/СОЧ
            if row_idx <= 15:
                print(f"\nСтрока {row_idx}: нет СОР/СОЧ | Тема: '{topic[:50]}'")

print("\n" + "="*60)
print(f"ВСЕГО НАЙДЕНО: {len(found_assessments)} записей")

for item in found_assessments:
    print(f"  {item['Класс']} | {item['Тип']} | №{item['Номер']} | {item['Дата']}")

input("\nНажмите Enter...")