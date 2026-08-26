import streamlit as st
import pandas as pd
from docx import Document
from io import BytesIO
import re

st.title("Тест СОР/СОЧ")

uploaded_file = st.file_uploader("Загрузите КТП", type=['docx'])

if uploaded_file:
    doc = Document(BytesIO(uploaded_file.read()))
    
    st.write(f"Таблиц: {len(doc.tables)}")
    
    results = []
    
    for table in doc.tables:
        for row_idx in range(1, len(table.rows)):
            row = table.rows[row_idx]
            
            try:
                topic = row.cells[2].text
                dates = row.cells[5].text
            except:
                continue
            
            # ПОИСК СОР
            sor_nums = re.findall(r'сор[^\d]*(\d+)', topic.lower())
            # ПОИСК СОЧ
            soch_nums = re.findall(r'соч[^\d]*(\d+)', topic.lower())
            
            if sor_nums or soch_nums:
                st.write(f"Строка {row_idx}: {topic.strip()[:50]}")
                st.write(f"  СОР: {sor_nums}, СОЧ: {soch_nums}")
                st.write(f"  Даты: {dates}")
                
                # Разбираем даты
                for line in dates.split('\n'):
                    m = re.match(r'(\d+[а-яА-Я]*)\s*-\s*([\d\.\-]+)', line.strip())
                    if m:
                        cls = m.group(1)
                        dt = m.group(2)
                        
                        for num in sor_nums:
                            results.append({'Класс': cls, 'Тип': 'СОР', 'Номер': num, 'Дата': dt})
                        for num in soch_nums:
                            results.append({'Класс': cls, 'Тип': 'СОЧ', 'Номер': num, 'Дата': dt})
    
    if results:
        st.success(f"Найдено: {len(results)}")
        df = pd.DataFrame(results)
        st.dataframe(df)
    else:
        st.warning("Ничего не найдено")
        # Показываем первые 10 тем
        for table in doc.tables:
            for row_idx in range(1, min(len(table.rows), 11)):
                try:
                    t = table.rows[row_idx].cells[2].text.strip()
                    if t:
                        st.write(f"{row_idx}: {t[:60]}")
                except:
                    pass