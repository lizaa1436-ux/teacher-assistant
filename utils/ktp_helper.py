from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime, timedelta, date
import pandas as pd
import json
import re
import io
import os


class KTPHelper:
    def __init__(self):
        self.calendar_helper = None

    def set_calendar_helper(self, calendar_helper):
        self.calendar_helper = calendar_helper

    def read_document(self, file_bytes, file_extension):
        """Читать документ Word (.docx или .doc)"""
        if file_extension.lower() == 'docx':
            return Document(io.BytesIO(file_bytes))
        elif file_extension.lower() == 'doc':
            try:
                return Document(io.BytesIO(file_bytes))
            except:
                raise ValueError("Файл .doc не поддерживается. Сохраните как .docx")
        else:
            raise ValueError(f"Неподдерживаемый формат: {file_extension}")

    def create_ktp_template(self, subject, class_name, total_hours, hours_per_week):
        """Создать шаблон КТП"""
        doc = Document()

        # Альбомная ориентация
        section = doc.sections[0]
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Cm(29.7)
        section.page_height = Cm(21.0)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)

        style = doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(10)

        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run('Календарно-тематическое планирование')
        run.bold = True
        run.font.size = Pt(14)

        p_info = doc.add_paragraph()
        p_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_info.add_run(f'{subject}, {class_name} класс')

        p_hours = doc.add_paragraph()
        p_hours.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_hours.add_run(f'Итого: {total_hours} часов, в неделю: {hours_per_week} часов')

        doc.add_paragraph()

        # Таблица 7 столбцов
        headers = ['№', 'Раздел', 'Темы', 'Цели обучения', 'Кол-во часов', 'Дата', 'Примечание']
        table = doc.add_table(rows=1, cols=7)
        table.style = 'Table Grid'

        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].runs[0].font.size = Pt(10)
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), 'D9E2F3')
            cell._tc.get_or_add_tcPr().append(shading)

        # Добавляем строки с четвертями
        lesson_number = 1
        quarter_hours = {1: total_hours // 4, 2: total_hours // 4, 3: total_hours // 4, 4: total_hours - (total_hours // 4) * 3}
        roman_map = {1: 'I', 2: 'II', 3: 'III', 4: 'IV'}

        for quarter in range(1, 5):
            # Заголовок четверти
            row = table.add_row()
            cell = row.cells[0]
            for i in range(1, 7):
                cell = cell.merge(row.cells[i])
            cell.text = f"{roman_map[quarter]} четверть"
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.paragraphs[0].runs[0].bold = True
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), 'E8EEF7')
            cell._tc.get_or_add_tcPr().append(shading)

            # Уроки
            for i in range(quarter_hours[quarter]):
                row = table.add_row()
                row.cells[0].text = str(lesson_number)
                row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                lesson_number += 1

            # Итог
            row = table.add_row()
            cell = row.cells[0]
            for i in range(1, 7):
                cell = cell.merge(row.cells[i])
            cell.text = f"Всего часов: {quarter_hours[quarter]}"
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            cell.paragraphs[0].runs[0].bold = True

        return doc, table

    def find_ktp_table(self, doc):
        """Найти таблицу КТП"""
        if not doc.tables:
            return None
        best_table = None
        best_score = 0
        for table in doc.tables:
            if len(table.rows) < 2:
                continue
            try:
                headers = [cell.text.strip().lower() for cell in table.rows[0].cells]
            except:
                continue
            score = 0
            if any('дат' in h or 'срок' in h for h in headers):
                score += 3
            if any('час' in h for h in headers):
                score += 3
            if any('тем' in h for h in headers):
                score += 2
            if len(table.columns) >= 7:
                score += 1
            if score > best_score:
                best_score = score
                best_table = table
        return best_table if best_score >= 6 else None

    def detect_columns(self, table):
        """Определить индексы столбцов"""
        col_indices = {}
        try:
            headers = [cell.text.strip().lower() for cell in table.rows[0].cells]
        except:
            headers = []
        for i, h in enumerate(headers):
            if h == '№' or 'номер' in h:
                col_indices['number'] = i
            elif 'раздел' in h:
                col_indices['section'] = i
            elif 'тем' in h:
                col_indices['topic'] = i
            elif 'цел' in h:
                col_indices['goals'] = i
            elif 'час' in h or 'кол-во' in h:
                col_indices['hours'] = i
            elif 'дат' in h or 'срок' in h:
                col_indices['date'] = i
            elif 'примеч' in h:
                col_indices['note'] = i
        # По умолчанию
        if 'hours' not in col_indices:
            col_indices['hours'] = 4
        if 'date' not in col_indices:
            col_indices['date'] = 5
        if 'number' not in col_indices:
            col_indices['number'] = 0
        if 'topic' not in col_indices:
            col_indices['topic'] = 2
        if 'goals' not in col_indices:
            col_indices['goals'] = 3
        return col_indices

    def check_cell_merge(self, cell):
        try:
            tc = cell._tc
            tcPr = tc.find(qn('w:tcPr'))
            if tcPr is not None:
                for child in tcPr:
                    tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
                    if tag in ['gridSpan', 'vMerge', 'hMerge']:
                        return True
        except:
            pass
        return False

    def check_row_has_merge(self, row):
        try:
            for cell in row.cells:
                if self.check_cell_merge(cell):
                    return True
        except:
            pass
        return False

    def parse_hours(self, text):
        if not text:
            return 0
        numbers = re.findall(r'\d+', text)
        return int(numbers[0]) if numbers else 0

    def _generate_dates_for_weekday(self, start_date, end_date, weekday):
        """Генерировать рабочие даты для дня недели"""
        working_dates = []
        current = start_date
        while current <= end_date:
            if current.weekday() == weekday:
                if self.calendar_helper:
                    if self.calendar_helper.is_working_day(current):
                        working_dates.append(current)
                else:
                    working_dates.append(current)
            current += timedelta(days=1)
        return working_dates

    def fill_dates_for_classes(self, doc, class_configs):
        """Заполнить даты СТРОГО в пределах четвертей"""
        ktp_table = self.find_ktp_table(doc)
        if ktp_table is None:
            raise ValueError("Не найдена таблица КТП")

        col_indices = self.detect_columns(ktp_table)

        # Генерируем даты для каждой четверти ОТДЕЛЬНО
        class_dates_by_quarter = {}
        for class_name, config in class_configs.items():
            class_dates_by_quarter[class_name] = {}
            for quarter in range(1, 5):
                q_start, q_end = self.calendar_helper.get_quarter_dates(quarter)
                if q_start and q_end:
                    quarter_dates = []
                    for lesson_info in config.get('lessons', []):
                        weekday = lesson_info.get('weekday', 0)
                        dates = self._generate_dates_for_weekday(q_start, q_end, weekday)
                        quarter_dates.extend(dates)
                    quarter_dates.sort()
                    class_dates_by_quarter[class_name][quarter] = quarter_dates

        date_indices = {}
        for class_name in class_configs:
            date_indices[class_name] = {1: 0, 2: 0, 3: 0, 4: 0}

        lessons_scheduled = 0
        compact_warnings = {}
        extra_lessons_warning = {}

        # НАХОДИМ ГРАНИЦЫ ЧЕТВЕРТЕЙ
        quarter_start_rows = {}

        for row_idx in range(1, len(ktp_table.rows)):
            row = ktp_table.rows[row_idx]
            all_text = ""
            for cell in row.cells:
                all_text += " " + cell.text.strip() + " "
            all_text_lower = all_text.lower()

            # Арабские цифры
            for q_num in range(1, 5):
                if q_num in quarter_start_rows:
                    continue
                if re.search(rf'{q_num}\s*четверть', all_text_lower):
                    quarter_start_rows[q_num] = row_idx

            # Римские цифры
            roman_map = {1: 'i', 2: 'ii', 3: 'iii', 4: 'iv'}
            for q_num, roman in roman_map.items():
                if q_num in quarter_start_rows:
                    continue
                if re.search(rf'\b{roman}\b\s*четверть', all_text_lower):
                    quarter_start_rows[q_num] = row_idx

        if 1 not in quarter_start_rows:
            quarter_start_rows[1] = 1

        # Диапазоны строк
        quarter_row_ranges = {}
        sorted_qs = sorted(quarter_start_rows.keys())

        for i, q in enumerate(sorted_qs):
            start_row = quarter_start_rows[q] + 1
            if i + 1 < len(sorted_qs):
                end_row = quarter_start_rows[sorted_qs[i + 1]] - 1
            else:
                end_row = len(ktp_table.rows) - 1
            quarter_row_ranges[q] = (start_row, end_row)

        # ЗАПОЛНЯЕМ ДАТЫ ПО ЧЕТВЕРТЯМ
        for quarter in range(1, 5):
            start_row, end_row = quarter_row_ranges.get(quarter, (0, -1))

            for class_name in class_configs:
                date_indices[class_name][quarter] = 0

            for row_idx in range(start_row, end_row + 1):
                if row_idx >= len(ktp_table.rows):
                    break

                row = ktp_table.rows[row_idx]

                # Пропускаем объединенные
                try:
                    has_merge = False
                    for cell in row.cells:
                        tc = cell._tc
                        tcPr = tc.find(qn('w:tcPr'))
                        if tcPr is not None:
                            for child in tcPr:
                                tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
                                if tag == 'gridSpan':
                                    has_merge = True
                                    break
                        if has_merge:
                            break
                    if has_merge:
                        continue
                except:
                    pass

                try:
                    topic_idx = col_indices.get('topic', 2)
                    hours_idx = col_indices.get('hours', 4)
                    date_idx = col_indices.get('date', 5)

                    topic_text = row.cells[topic_idx].text.strip() if topic_idx < len(row.cells) else ""
                    hours_text = row.cells[hours_idx].text.strip() if hours_idx < len(row.cells) else ""
                    date_cell = row.cells[date_idx] if date_idx < len(row.cells) else None
                except:
                    continue

                if not topic_text:
                    continue

                hours = self.parse_hours(hours_text)

                if hours == 0:
                    is_sor_soch = bool(re.search(r'СО[рРчЧ]\s*№?\s*\d*', topic_text, re.IGNORECASE))
                    if is_sor_soch:
                        hours = 1

                if hours > 0:
                    date_strings = []

                    for class_name in class_configs:
                        idx = date_indices[class_name][quarter]
                        dates_list = class_dates_by_quarter[class_name].get(quarter, [])
                        lesson_dates = []

                        for h in range(hours):
                            if idx + h < len(dates_list):
                                lesson_dates.append(dates_list[idx + h].strftime('%d.%m'))
                            else:
                                lesson_dates.append('-')
                                if class_name not in compact_warnings:
                                    compact_warnings[class_name] = {}
                                if quarter not in compact_warnings[class_name]:
                                    compact_warnings[class_name][quarter] = 0
                                compact_warnings[class_name][quarter] += 1

                        if len(lesson_dates) == 1:
                            date_strings.append(f"{class_name} - {lesson_dates[0]}")
                        else:
                            date_strings.append(f"{class_name} - {'/'.join(lesson_dates)}")

                        date_indices[class_name][quarter] += hours

                    if date_cell is not None:
                        try:
                            for paragraph in date_cell.paragraphs:
                                for run in paragraph.runs:
                                    run.text = ""
                            date_cell.text = '\n'.join(date_strings)
                            for paragraph in date_cell.paragraphs:
                                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                for run in paragraph.runs:
                                    run.font.size = Pt(9)
                            lessons_scheduled += 1
                        except:
                            pass

        # Свободные даты
        for class_name in class_configs:
            for quarter in range(1, 5):
                dates_list = class_dates_by_quarter[class_name].get(quarter, [])
                used = date_indices[class_name][quarter]
                remaining = len(dates_list) - used
                if remaining > 0:
                    free_dates = []
                    for i in range(used, len(dates_list)):
                        free_dates.append(dates_list[i].strftime('%d.%m'))
                    if class_name not in extra_lessons_warning:
                        extra_lessons_warning[class_name] = {}
                    extra_lessons_warning[class_name][quarter] = {
                        'count': remaining,
                        'dates': free_dates
                    }

        debug_info = {
            'lessons_scheduled': lessons_scheduled,
            'col_indices': col_indices,
            'quarter_start_rows': quarter_start_rows,
            'quarter_ranges': quarter_row_ranges,
            'class_dates_count': {
                cls: {q: len(dates) for q, dates in quarters.items()}
                for cls, quarters in class_dates_by_quarter.items()
            }
        }

        return doc, lessons_scheduled, compact_warnings, extra_lessons_warning, debug_info

    def find_sor_soch(self, doc):
        """Найти СОР и СОЧ в документе"""
        ktp_table = self.find_ktp_table(doc)
        if ktp_table is None:
            return []

        col_indices = self.detect_columns(ktp_table)
        sor_soch_items = []

        patterns = [
            (r'СОр\s*№?\s*(\d+)', 'СОР'),
            (r'СОч\s*№?\s*(\d+)', 'СОЧ'),
            (r'[СC]ор\s*№?\s*(\d+)', 'СОР'),
            (r'[СC]оч\s*№?\s*(\d+)', 'СОЧ'),
            (r'СОР\s*№?\s*(\d+)', 'СОР'),
            (r'СОЧ\s*№?\s*(\d+)', 'СОЧ'),
        ]

        for row_idx in range(1, len(ktp_table.rows)):
            row = ktp_table.rows[row_idx]

            if self.check_row_has_merge(row):
                continue

            try:
                topic = row.cells[col_indices.get('topic', 2)].text.strip()
                date_text = row.cells[col_indices.get('date', 5)].text.strip()
            except:
                continue

            for pattern, sor_type in patterns:
                match = re.search(pattern, topic, re.IGNORECASE)
                if match:
                    number = match.group(1)
                    class_dates = self._parse_class_dates(date_text)
                    sor_soch_items.append({
                        'row': row_idx,
                        'topic': topic,
                        'type': sor_type,
                        'number': number,
                        'class_dates': class_dates
                    })
                    break

        return sor_soch_items

    def _parse_class_dates(self, date_text):
        """Разобрать даты по классам"""
        class_dates = {}
        if not date_text:
            return class_dates
        lines = date_text.strip().split('\n')
        for line in lines:
            line = line.strip()
            if not line:
                continue
            match = re.match(r'([\d\wа-яА-Я]+)\s*[-–—]\s*(.+)', line)
            if match:
                class_name = match.group(1).strip()
                dates_str = match.group(2).strip()
                dates = [d.strip() for d in dates_str.split('/') if d.strip()]
                class_dates[class_name] = dates
        return class_dates

    def create_sor_soch_schedule(self, doc):
        """Создать график СОР и СОЧ"""
        sor_soch_items = self.find_sor_soch(doc)

        if not sor_soch_items:
            return None, "Не найдены СОР/СОЧ в документе"

        schedule = []

        for item in sor_soch_items:
            for class_name, dates in item['class_dates'].items():
                schedule.append({
                    'Класс': class_name,
                    'Тип': item['type'],
                    'Номер': item['number'],
                    'Тема': item['topic'][:50],
                    'Дата': ', '.join(dates) if dates else 'Не указана'
                })

        return schedule, None


def validate_ktp_structure(doc):
    """Проверить структуру КТП"""
    try:
        ktp_helper = KTPHelper()
        ktp_table = ktp_helper.find_ktp_table(doc)
        if ktp_table is None:
            return False, "Не найдена таблица КТП"
        headers = [cell.text.strip() for cell in ktp_table.rows[0].cells]
        if len(headers) != 7:
            return False, f"Должно быть 7 столбцов, найдено: {len(headers)}"
        return True, f"Структура верная. Строк: {len(ktp_table.rows) - 1}"
    except Exception as e:
        return False, f"Ошибка: {e}"