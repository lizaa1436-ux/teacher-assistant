from docx import Document
import re
import os
import sys

def check_row_merged(row):
    """Проверить объединенные ячейки"""
    try:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tcPr')
            if tcPr is not None:
                for child in tcPr:
                    tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
                    if tag in ['gridSpan', 'vMerge']:
                        return True
    except:
        pass
    return False


def find_sor_soch(text):
    """Простой поиск СОР/СОЧ"""
    if not text:
        return []
    
    text_lower = text.lower()
    found = []
    
    if 'сор' in text_lower:
        numbers = re.findall(r'сор[^\d]*(\d+)', text_lower)
        found.append(f"СОР: {numbers}")
    
    if 'соч' in text_lower:
        numbers = re.findall(r'соч[^\d]*(\d+)', text_lower)
        found.append(f"СОЧ: {numbers}")
    
    return found


# Запросить путь к файлу
file_path = input("Введите путь к файлу КТП: ").strip().strip('"')

# Проверяем существование файла
if not os.path.exists(file_path):
    print(f"❌ Файл не найден: {file_path}")
    print("Попробуйте:")
    print("1. Перетащить файл в окно терминала")
    print("2. Или ввести путь без кавычек")
    sys.exit(1)

print(f"Файл существует: {file_path}")
print(f"Размер: {os.path.getsize(file_path)} байт")
print(f"Расширение: {os.path.splitext(file_path)[1]}")

try:
    doc = Document(file_path)
    print(f"✅ Файл открыт успешно")
    print(f"Таблиц в документе: {len(doc.tables)}")
    
    # Проходим по всем таблицам
    for table_idx, table in enumerate(doc.tables):
        print(f"\n{'='*60}")
        print(f"ТАБЛИЦА {table_idx}")
        print(f"Строк: {len(table.rows)}")
        
        # Заголовки
        if len(table.rows) > 0:
            print(f"\n--- ЗАГОЛОВКИ ---")
            for i, cell in enumerate(table.rows[0].cells):
                print(f"  Столбец {i}: '{cell.text.strip()}'")
        
        # Первые 20 строк
        print(f"\n--- СТРОКИ ---")
        for row_idx in range(1, min(len(table.rows), 21)):
            row = table.rows[row_idx]
            merged = check_row_merged(row)
            
            cells_text = []
            for i, cell in enumerate(row.cells):
                text = cell.text.strip()
                cells_text.append(f"[{i}]='{text[:50]}'")
            
            full_text = ' '.join([cell.text for cell in row.cells])
            sor_soch = find_sor_soch(full_text)
            
            merge_str = "ОБЪЕДИНЕНА" if merged else ""
            
            print(f"\nСтрока {row_idx} {merge_str}:")
            for ct in cells_text:
                print(f"  {ct}")
            if sor_soch:
                print(f"  ⭐ НАЙДЕНО: {', '.join(sor_soch)}")
    
except Exception as e:
    print(f"\n❌ Ошибка: {e}")
    import traceback
    traceback.print_exc()

input("\nНажмите Enter для выхода...")